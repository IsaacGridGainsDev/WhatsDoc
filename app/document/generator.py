import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

def generate_document(content, output_path, format_type, branding=None):
    """
    Generate a document from markdown content
    
    Args:
        content (str): Markdown content to convert
        output_path (str): Path to save the document
        format_type (str): 'pdf' or 'docx'
        branding (dict, optional): Branding information for header/footer
        
    Returns:
        str: Path to the generated document
    """
    if format_type.lower() == "pdf":
        return generate_pdf(content, output_path, branding)
    elif format_type.lower() == "docx":
        return generate_docx(content, output_path, branding)
    else:
        raise ValueError(f"Unsupported format type: {format_type}")

def generate_pdf(content, output_path, branding=None):
    """
    Generate a PDF document from markdown content
    
    Args:
        content (str): Markdown content to convert
        output_path (str): Path to save the PDF
        branding (dict, optional): Branding information for header/footer
        
    Returns:
        str: Path to the generated PDF
    """
    # Parse markdown content
    parsed_content = parse_markdown(content)
    
    # Create PDF
    pdf = FPDF()
    pdf.add_page()
    
    # Set default font
    pdf.set_font("Arial", size=12)
    
    # Add branding if provided
    if branding:
        add_branding_to_pdf(pdf, branding)
    
    # Add content to PDF
    for item in parsed_content:
        if item["type"] == "heading1":
            pdf.set_font("Arial", "B", 18)
            pdf.cell(0, 10, item["content"], ln=True)
            pdf.ln(5)
        elif item["type"] == "heading2":
            pdf.set_font("Arial", "B", 16)
            pdf.cell(0, 10, item["content"], ln=True)
            pdf.ln(5)
        elif item["type"] == "heading3":
            pdf.set_font("Arial", "B", 14)
            pdf.cell(0, 10, item["content"], ln=True)
            pdf.ln(5)
        elif item["type"] == "paragraph":
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, item["content"])
            pdf.ln(5)
        elif item["type"] == "bullet":
            pdf.set_font("Arial", size=12)
            pdf.cell(10, 10, "•", ln=0)
            pdf.multi_cell(0, 10, item["content"])
    
    # Save PDF
    pdf.output(output_path)
    
    return output_path

def generate_docx(content, output_path, branding=None):
    """
    Generate a DOCX document from markdown content
    
    Args:
        content (str): Markdown content to convert
        output_path (str): Path to save the DOCX
        branding (dict, optional): Branding information for header/footer
        
    Returns:
        str: Path to the generated DOCX
    """
    # Parse markdown content
    parsed_content = parse_markdown(content)
    
    # Create DOCX
    doc = Document()
    
    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(12)
    
    # Add branding if provided
    if branding:
        add_branding_to_docx(doc, branding)
    
    # Add content to DOCX
    for item in parsed_content:
        if item["type"] == "heading1":
            heading = doc.add_heading(item["content"], level=1)
            heading.style.font.size = Pt(18)
        elif item["type"] == "heading2":
            heading = doc.add_heading(item["content"], level=2)
            heading.style.font.size = Pt(16)
        elif item["type"] == "heading3":
            heading = doc.add_heading(item["content"], level=3)
            heading.style.font.size = Pt(14)
        elif item["type"] == "paragraph":
            doc.add_paragraph(item["content"])
        elif item["type"] == "bullet":
            doc.add_paragraph(item["content"], style="List Bullet")
    
    # Save DOCX
    doc.save(output_path)
    
    return output_path

def parse_markdown(content):
    """
    Parse markdown content into structured format
    
    Args:
        content (str): Markdown content
        
    Returns:
        list: List of parsed content items
    """
    parsed_content = []
    lines = content.split("\n")
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:  # Skip empty lines
            i += 1
            continue
        
        # Check for headings
        if line.startswith("# "):
            parsed_content.append({"type": "heading1", "content": line[2:]})
        elif line.startswith("## "):
            parsed_content.append({"type": "heading2", "content": line[3:]})
        elif line.startswith("### "):
            parsed_content.append({"type": "heading3", "content": line[4:]})
        # Check for bullet points
        elif line.startswith("- ") or line.startswith("* "):
            parsed_content.append({"type": "bullet", "content": line[2:]})
        # Regular paragraph
        else:
            # Combine multi-line paragraphs
            paragraph = line
            j = i + 1
            while j < len(lines) and lines[j].strip() and not lines[j].strip().startswith(("#", "-", "*")):
                paragraph += " " + lines[j].strip()
                j += 1
            i = j - 1  # Adjust index to account for combined lines
            
            parsed_content.append({"type": "paragraph", "content": paragraph})
        
        i += 1
    
    return parsed_content

def add_branding_to_pdf(pdf, branding):
    """
    Add branding elements to PDF
    
    Args:
        pdf (FPDF): PDF document
        branding (dict): Branding information
    """
    # Add header
    if "header" in branding:
        pdf.set_font("Arial", "B", 10)
        pdf.cell(0, 10, branding["header"], 0, 1, "R")
        pdf.ln(5)
    
    # Add logo if provided
    if "logo_path" in branding and os.path.exists(branding["logo_path"]):
        pdf.image(branding["logo_path"], x=10, y=10, w=30)
    
    # Set footer
    if "footer" in branding:
        pdf.set_y(-15)
        pdf.set_font("Arial", "I", 8)
        pdf.cell(0, 10, branding["footer"], 0, 0, "C")

def add_branding_to_docx(doc, branding):
    """
    Add branding elements to DOCX
    
    Args:
        doc (Document): DOCX document
        branding (dict): Branding information
    """
    # Add header
    if "header" in branding:
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = branding["header"]
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add logo if provided
    if "logo_path" in branding and os.path.exists(branding["logo_path"]):
        doc.add_picture(branding["logo_path"], width=Inches(1))
    
    # Set footer
    if "footer" in branding:
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = branding["footer"]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER